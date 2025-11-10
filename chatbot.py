"""
rag_chatbot_strict.py
---------------------
Retrieval-Augmented Generation (RAG) chatbot that ONLY answers from the provided documents.

Features:
- Load data files from ./data (supports .txt, .pdf, .csv, .docx)
- Build a Chroma vector store of text chunks
- Expose a FastAPI endpoint /chat that:
    - Retrieves relevant chunks
    - Sends them with a strict system prompt to the LLM that enforces:
        * Answer ONLY using provided context
        * If information is missing, reply: "I don't know based on the provided documents."
- Endpoint /build to create or refresh the index
- Returns sources used for the answer

Requirements (put into requirements.txt):
    fastapi
    uvicorn
    langchain
    langchain-openai
    chromadb
    openai
    tiktoken
    pandas
    PyPDF2
    python-docx
    pydantic

Usage:
  1) Put documents into ./data (txt, pdf, csv, docx)
  2) export OPENAI_API_KEY="sk-..."
  3) Build index (one-time or when data changes):
       python rag_chatbot_strict.py build
  4) Run server:
       uvicorn rag_chatbot_strict:app --reload
  5) Ask:
       POST http://127.0.0.1:8000/chat  with JSON {"question": "your question"}
"""

import os
import json
import glob
from typing import List, Dict, Any
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from pathlib import Path

# LangChain + Chroma + OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LCDocument
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.vectorstores import Chroma
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate

# Document loading
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# Utility
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("rag_strict")

# CONFIG
DATA_DIR = "data"                       # put your files here
PERSIST_DIR = "chroma_store"            # where to persist embeddings
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 100
RETRIEVE_K = 4                          # number of chunks to retrieve
OPENAI_MODEL = "gpt-4o"                 # change if you prefer another model (must be available)
LLM_TEMPERATURE = 0.0                   # deterministic answers
MAX_CONTEXT_TOKENS = 3000               # safety guard for context size

# System prompt enforces strict grounding and the exact fallback sentence.
SYSTEM_PROMPT = """You are a helpful assistant that MUST answer questions using ONLY the information in the provided CONTEXT.
Follow these rules exactly:
1) Use only facts present in the CONTEXT. Do NOT use external knowledge, web searches, or assumptions.
2) If the CONTEXT contains the answer, produce a concise answer and include short citations in square brackets using the source filenames, for example: (source1.pdf).
3) If the CONTEXT does NOT contain enough information to answer, reply exactly: "I don't know based on the provided documents."
4) Do not hallucinate. If uncertain, use the exact fallback sentence above.
5) Keep answers short and factual.
"""

# Build prompt template (we will insert the retrieved context and the user question)
PROMPT_TEMPLATE = """{system_instructions}

CONTEXT:
{context}

QUESTION:
{question}

INSTRUCTIONS: Answer the question using ONLY the context above. If the answer cannot be found in the context, reply exactly: "I don't know based on the provided documents."
If you answer, include which source(s) you used in parentheses at the end, e.g. (report.pdf).
"""

# FastAPI app
app = FastAPI(title="RAG Chatbot (Strictly Grounded)")

# Pydantic models
class BuildRequest(BaseModel):
    data_dir: str = DATA_DIR  # optional override

class ChatRequest(BaseModel):
    question: str

# -------------------------
# Document loaders
# -------------------------
def load_txt(path: str) -> List[LCDocument]:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [LCDocument(page_content=text, metadata={"source": Path(path).name})]

def load_pdf(path: str) -> List[LCDocument]:
    text = ""
    try:
        reader = PdfReader(path)
        for p in reader.pages:
            txt = p.extract_text()
            if txt:
                text += txt + "\n"
    except Exception as e:
        logger.warning(f"PDF load failed for {path}: {e}")
    return [LCDocument(page_content=text, metadata={"source": Path(path).name})]

def load_docx(path: str) -> List[LCDocument]:
    text = []
    try:
        doc = DocxDocument(path)
        for para in doc.paragraphs:
            text.append(para.text)
    except Exception as e:
        logger.warning(f"DOCX load failed for {path}: {e}")
    return [LCDocument(page_content="\n".join(text), metadata={"source": Path(path).name})]

def load_csv(path: str, text_columns: List[str] = None) -> List[LCDocument]:
    """
    Convert each row into a text document by concatenating chosen columns.
    If text_columns is None, use all columns.
    """
    df = pd.read_csv(path)
    if text_columns is None:
        cols = df.columns.tolist()
    else:
        cols = text_columns
    docs = []
    for _, row in df.iterrows():
        parts = []
        for c in cols:
            parts.append(f"{c}: {row.get(c, '')}")
        txt = " \n".join(parts)
        docs.append(LCDocument(page_content=txt, metadata={"source": Path(path).name}))
    return docs

def load_documents_from_folder(folder: str = DATA_DIR) -> List[LCDocument]:
    """Load supported files from folder into a list of LangChain Documents."""
    docs = []
    folder = Path(folder)
    if not folder.exists():
        raise FileNotFoundError(f"Data folder {folder} does not exist")
    supported = list(folder.glob("*"))
    for f in supported:
        ext = f.suffix.lower()
        try:
            if ext == ".txt":
                docs.extend(load_txt(str(f)))
            elif ext == ".pdf":
                docs.extend(load_pdf(str(f)))
            elif ext == ".docx":
                docs.extend(load_docx(str(f)))
            elif ext == ".csv":
                docs.extend(load_csv(str(f)))
            else:
                logger.info(f"Skipping unsupported file {f.name}")
        except Exception as e:
            logger.exception(f"Error loading {f.name}: {e}")
    logger.info(f"Loaded {len(docs)} raw documents from {folder}")
    return docs

# -------------------------
# Index / Vector store build
# -------------------------
def build_vector_store(data_dir: str = DATA_DIR, persist_dir: str = PERSIST_DIR, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
    """
    Loads documents, chunk-splits them, creates embeddings, and persists to Chroma.
    """
    docs = load_documents_from_folder(data_dir)
    if not docs:
        raise HTTPException(status_code=400, detail="No documents found to index.")

    # Split into chunks
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = splitter.split_documents(docs)
    logger.info(f"Split into {len(chunks)} chunks")

    # Create embeddings and persist
    embeddings = OpenAIEmbeddings()
    vectordb = Chroma.from_documents(chunks, embedding=embeddings, persist_directory=persist_dir)
    vectordb.persist()
    logger.info(f"Saved Chroma at {persist_dir}")
    return {"status": "ok", "chunks_indexed": len(chunks), "persist_dir": persist_dir}

# -------------------------
# Retrieval + Answering
# -------------------------
def load_vectordb(persist_dir: str = PERSIST_DIR):
    embeddings = OpenAIEmbeddings()
    if not Path(persist_dir).exists():
        raise FileNotFoundError(f"Vector store at {persist_dir} not found. Build it first.")
    vectordb = Chroma(persist_directory=persist_dir, embedding_function=embeddings)
    return vectordb

def create_prompt(context: str, question: str) -> str:
    prompt = PROMPT_TEMPLATE.format(system_instructions=SYSTEM_PROMPT, context=context, question=question)
    return prompt

def truncate_context_by_tokens(context: str, max_tokens: int = MAX_CONTEXT_TOKENS) -> str:
    """
    Very simple guard — truncate by characters since token counting library may vary.
    This aims to keep context under a rough token budget.
    """
    if len(context) <= max_tokens * 4:  # rough 4 chars per token
        return context
    return context[: max_tokens * 4]

def retrieve_and_answer(question: str, k: int = RETRIEVE_K, persist_dir: str = PERSIST_DIR) -> Dict[str, Any]:
    """
    1) Load vector store
    2) Retrieve top-k chunks
    3) Build a context string combining chunk text + source markers
    4) Ask the LLM with a strict prompt
    5) Return the answer and sources
    """
    vectordb = load_vectordb(persist_dir)
    retriever = vectordb.as_retriever(search_kwargs={"k": k})

    # get documents
    docs = retriever.get_relevant_documents(question)
    if not docs:
        return {"answer": "I don't know based on the provided documents.", "sources": []}

    # Build context using retrieved docs. Include filename metadata.
    # We'll keep a running set of source filenames to cite.
    context_parts = []
    sources = []
    for i, d in enumerate(docs):
        src = d.metadata.get("source", f"chunk_{i}")
        # Deduplicate sources
        if src not in sources:
            sources.append(src)
        # Add clear separators and source tags so LLM can cite properly
        context_parts.append(f"=== SOURCE: {src} ===\n{d.page_content}\n")

    context = "\n\n".join(context_parts)
    context = truncate_context_by_tokens(context)

    # Create LLM prompt
    prompt = create_prompt(context=context, question=question)

    # Call LLM
    llm = ChatOpenAI(model_name=OPENAI_MODEL, temperature=LLM_TEMPERATURE)
    # Use an LLMChain with a simple template to pass system prompt + user prompt
    prompt_template = PromptTemplate(input_variables=["system_instructions", "context", "question"],
                                     template=PROMPT_TEMPLATE)
    chain = LLMChain(llm=llm, prompt=prompt_template)

    # Run chain
    response = chain.run({"system_instructions": SYSTEM_PROMPT, "context": context, "question": question})
    # The LLM should either answer with content + (source.pdf) or exactly the fallback.
    # We still sanity-check: if the response contains fallback phrase, return fallback.
    fallback = "I don't know based on the provided documents."
    if fallback.lower() in response.strip().lower():
        return {"answer": fallback, "sources": []}
    # Otherwise attach the sources list (we expect model to include parentheses; still return sources separately)
    return {"answer": response.strip(), "sources": sources}

# -------------------------
# FastAPI endpoints
# -------------------------
@app.post("/build", summary="Build or rebuild the vector store from files in data folder")
def build_endpoint(req: BuildRequest):
    try:
        result = build_vector_store(data_dir=req.data_dir)
        return result
    except Exception as e:
        logger.exception("Build failed")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/chat", summary="Ask a question grounded in provided documents")
def chat_endpoint(req: ChatRequest):
    q = req.question.strip()
    if not q:
        raise HTTPException(status_code=400, detail="Question cannot be empty")
    try:
        result = retrieve_and_answer(q)
        return result
    except FileNotFoundError as fnf:
        raise HTTPException(status_code=400, detail=str(fnf))
    except Exception as e:
        logger.exception("Chat failed")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
def health():
    return {"status": "ok", "data_dir_exists": os.path.exists(DATA_DIR), "vectorstore_exists": os.path.exists(PERSIST_DIR)}

# -------------------------
# CLI: allow quick build from command line
# -------------------------
if __name__ == "__main__":
    import argparse
    p = argparse.ArgumentParser()
    sub = p.add_subparsers(dest="cmd")

    b = sub.add_parser("build", help="Build vector store from data folder")
    b.add_argument("--data_dir", default=DATA_DIR)
    b.add_argument("--persist_dir", default=PERSIST_DIR)
    b.add_argument("--chunk_size", default=CHUNK_SIZE, type=int)
    b.add_argument("--chunk_overlap", default=CHUNK_OVERLAP, type=int)

    q = sub.add_parser("query", help="Query the built index (local single-shot)")
    q.add_argument("question", type=str)
    q.add_argument("--k", default=RETRIEVE_K, type=int)

    args = p.parse_args()
    if args.cmd == "build":
        print("Building index...")
        out = build_vector_store(data_dir=args.data_dir, persist_dir=args.persist_dir, chunk_size=args.chunk_size, chunk_overlap=args.chunk_overlap)
        print("Done:", out)
    elif args.cmd == "query":
        print("Querying...")
        res = retrieve_and_answer(args.question, k=args.k)
        print(json.dumps(res, indent=2))
    else:
        print("Run uvicorn to serve API, or use 'build'/'query' commands.")








# Example quick run

# Put files into ./data: spec.pdf, faq.txt, records.csv

# Build index:

# python rag_chatbot_strict.py build


# Run server:

# uvicorn rag_chatbot_strict:app --reload


# Chat:

# curl -X POST "http://127.0.0.1:8000/chat" -H "Content-Type: application/json" -d '{"question":"What is the refund policy?"}'


# If the documents contain the refund policy, the bot will answer using the text and cite sources; otherwise it will reply:

# "I don't know based on the provided documents."